import os

# Set up matplotlib configuration directory and headless backend for Vercel Serverless environment
os.environ["MPLCONFIGDIR"] = "/tmp"
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from flask import Flask, render_template, request, jsonify, session, send_file
from google import genai
from google.genai import types as genai_types
import json
import re
import zipfile
import io
import pandas as pd

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "research_mentor_secret_key_192837")

# Vercel Configuration Settings (For Vercel Pro instances allowing up to 60s processing timeouts)
# NOTE FOR VERCEL HOBBY USERS: Serverless functions on Vercel Hobby profiles timeout after 10 seconds.
# Sequential AI operations executed during multi-section generation will exceed Hobby limits.
maxDuration = 60 

def get_client():
    key = os.getenv("GEMINI_API_KEY", "").strip()
    if not key:
        return None
    try:
        return genai.Client(api_key=key)
    except Exception:
        return None

def gemini_call(prompt: str, temperature: float = 0.4) -> str:
    client = get_client()
    if not client:
        return "ERROR: No API key configured. Configure GEMINI_API_KEY in your Vercel Environment Variables."
    try:
        cfg = genai_types.GenerateContentConfig(
            temperature=temperature,
            max_output_tokens=2048,
        )
        resp = client.models.generate_content(
            model="gemini-2.0-flash-lite",
            contents=prompt,
            config=cfg,
        )
        return resp.text.strip()
    except Exception as e:
        err = str(e)
        if "quota" in err.lower() or "rate" in err.lower():
            return "ERROR: API rate limit reached. Wait a moment and try again."
        if "invalid" in err.lower() or "api key" in err.lower():
            return "ERROR: Invalid API key. Check your environment configuration."
        return f"ERROR: {err[:200]}"

def is_error(text: str) -> bool:
    return text.startswith("ERROR:")

def extract_pdf_metadata(file_bytes: bytes) -> dict:
    if not PYMUPDF_AVAILABLE:
        return {"title": "PDF upload (PyMuPDF not installed)", "authors": "Unknown", "abstract": "", "keywords": ""}
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc[:4]:
            text += page.get_text()
        doc.close()
        text = text[:4000]

        title = ""
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        if lines:
            title = lines[0][:120]

        abstract = ""
        abs_match = re.search(r'(?i)abstract[\s\n:]+(.{100,1200}?)(?=\n\n|\nintroduction|\nkeyword)', text, re.DOTALL)
        if abs_match:
            abstract = abs_match.group(1).replace("\n", " ").strip()

        keywords = ""
        kw_match = re.search(r'(?i)keywords?[\s\n:—]+([^\n]{20,300})', text)
        if kw_match:
            keywords = kw_match.group(1).strip()

        author_match = re.search(r'(?i)(author|by)[\s:]+([^\n]{5,150})', text)
        authors = author_match.group(2).strip() if author_match else "See paper"

        return {"title": title or "Untitled", "authors": authors, "abstract": abstract, "keywords": keywords, "raw_text": text}
    except Exception as e:
        return {"title": "Parse error", "authors": "", "abstract": "", "keywords": "", "error": str(e)}

def init_session_state():
    if "project_context" not in session:
        session["project_context"] = {}
    if "uploaded_references" not in session:
        session["uploaded_references"] = []
    if "selected_topic" not in session:
        session["selected_topic"] = ""
    if "research_gap" not in session:
        session["research_gap"] = {}
    if "thesis" not in session:
        session["thesis"] = {}
    if "outline" not in session:
        session["outline"] = []
    if "paper_sections" not in session:
        session["paper_sections"] = {}
    if "citations" not in session:
        session["citations"] = {"verified": [], "unverified": []}
    if "review_results" not in session:
        session["review_results"] = {}
    if "generated_latex" not in session:
        session["generated_latex"] = {}
    if "experiment_data" not in session:
        session["experiment_data"] = {}
    if "current_step" not in session:
        session["current_step"] = 1

@app.route("/")
def index():
    init_session_state()
    return render_template("index.html")

@app.route("/api/state", methods=["GET"])
def get_state():
    init_session_state()
    return jsonify({
        "project_context": session.get("project_context"),
        "uploaded_references": session.get("uploaded_references"),
        "selected_topic": session.get("selected_topic"),
        "research_gap": session.get("research_gap"),
        "thesis": session.get("thesis"),
        "outline": session.get("outline"),
        "paper_sections": session.get("paper_sections"),
        "citations": session.get("citations"),
        "review_results": session.get("review_results"),
        "generated_latex": session.get("generated_latex"),
        "experiment_data": session.get("experiment_data"),
        "current_step": session.get("current_step")
    })

@app.route("/api/step", methods=["POST"])
def update_step():
    data = request.get_json() or {}
    session["current_step"] = data.get("step", 1)
    return jsonify({"success": True})

@app.route("/api/setup/save", methods=["POST"])
def save_setup():
    data = request.get_json() or {}
    session["project_context"] = {
        "title": data.get("title", ""),
        "subject": data.get("subject", ""),
        "domain": data.get("domain", ""),
        "level": data.get("level", "BTech"),
        "citation_style": data.get("citation_style", "IEEE"),
        "journal": data.get("journal", "IEEE"),
        "length": data.get("length", "6-8 pages"),
        "deadline": data.get("deadline", ""),
        "student_names": data.get("student_names", ""),
        "guide": data.get("guide", ""),
        "college": data.get("college", ""),
        "prompt": data.get("prompt", "")
    }
    session["current_step"] = 2
    return jsonify({"success": True})

@app.route("/api/references/upload", methods=["POST"])
def upload_reference():
    if "pdf_file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file = request.files["pdf_file"]
    if file.filename == "":
        return jsonify({"error": "Empty filename"}), 400

    refs = session.get("uploaded_references", [])
    existing_names = {r["filename"] for r in refs}

    if file.filename not in existing_names:
        file_bytes = file.read()
        meta = extract_pdf_metadata(file_bytes)
        meta["filename"] = file.filename
        refs.append(meta)
        session["uploaded_references"] = refs
        return jsonify({"success": True, "reference": meta})
    return jsonify({"success": False, "message": "File already exists"})

@app.route("/api/references/remove", methods=["POST"])
def remove_reference():
    data = request.get_json() or {}
    index = data.get("index", -1)
    refs = session.get("uploaded_references", [])
    if 0 <= index < len(refs):
        refs.pop(index)
        session["uploaded_references"] = refs
        return jsonify({"success": True})
    return jsonify({"success": False, "error": "Invalid index"}), 400

@app.route("/api/topics/generate", methods=["POST"])
def api_generate_topics():
    ctx = session.get("project_context", {})
    if not ctx.get("subject"):
        return jsonify({"error": "Complete Assignment Setup first."}), 400

    prompt = f"""You are a research advisor. Suggest 5 research topics for a {ctx.get('level','BTech')} student.
Subject: {ctx.get('subject','')}
Domain: {ctx.get('domain','')}

Return ONLY a JSON array. Each item must have: title, difficulty (Easy/Medium/Hard), novelty (Low/Medium/High), potential (Low/Medium/High), rationale (one sentence).
No markdown. No explanation. Just the JSON array."""
    
    raw = gemini_call(prompt, 0.6)
    if is_error(raw):
        return jsonify({"error": raw}), 500
    try:
        clean = re.sub(r'```json|```', '', raw).strip()
        parsed = json.loads(clean)
        return jsonify({"topics": parsed})
    except Exception:
        fallback = [{"title": t.strip(), "difficulty": "Medium", "novelty": "Medium", "potential": "Medium", "rationale": ""} 
                    for t in raw.split("\n") if t.strip() and len(t.strip()) > 10][:5]
        return jsonify({"topics": fallback})

@app.route("/api/topics/select", methods=["POST"])
def select_topic():
    data = request.get_json() or {}
    topic = data.get("topic", "")
    session["selected_topic"] = topic
    session["current_step"] = 4
    return jsonify({"success": True})

@app.route("/api/gap/analyze", methods=["POST"])
def api_analyze_gap():
    topic = session.get("selected_topic", "")
    if not topic:
        return jsonify({"error": "Select a topic first."}), 400

    refs = session.get("uploaded_references", [])
    refs_text = " | ".join([f"{r.get('title','')} — {r.get('abstract','')[:300]}" for r in refs])

    prompt = f"""You are a research analyst. Based on the topic and reference summaries below, identify the research gap.
Topic: {topic}
References (summaries): {refs_text[:2000] if refs_text else 'No references uploaded.'}

Return ONLY a JSON object with keys: current_trends (list of 3 strings), limitations (list of 3 strings), research_gaps (list of 3 strings), future_opportunities (list of 3 strings).
No markdown. No explanation."""

    raw = gemini_call(prompt)
    if is_error(raw):
        return jsonify({"error": raw}), 500
    try:
        clean = re.sub(r'```json|```', '', raw).strip()
        parsed = json.loads(clean)
        session["research_gap"] = parsed
        return jsonify({"gap": parsed})
    except Exception:
        fallback = {"current_trends": [], "limitations": [], "research_gaps": [raw[:300]], "future_opportunities": []}
        session["research_gap"] = fallback
        return jsonify({"gap": fallback})

@app.route("/api/thesis/generate", methods=["POST"])
def api_generate_thesis():
    topic = session.get("selected_topic", "")
    ctx = session.get("project_context", {})
    if not topic:
        return jsonify({"error": "Select a topic first."}), 400

    gap_data = session.get("research_gap", {})
    gap_str = " | ".join(gap_data.get("research_gaps", ["No gap identified yet."]))

    prompt = f"""You are a thesis advisor helping a student build a clear thesis statement.
Topic: {topic}
Research Gap: {gap_str}
Subject: {ctx.get('subject','')}

Return ONLY a JSON object with keys: problem_statement, research_question, hypothesis, objectives (list of 3-4 strings), expected_contribution.
Keep each field concise. No markdown. No explanation."""

    raw = gemini_call(prompt)
    if is_error(raw):
        return jsonify({"error": raw}), 500
    try:
        clean = re.sub(r'```json|```', '', raw).strip()
        parsed = json.loads(clean)
        session["thesis"] = parsed
        return jsonify({"thesis": parsed})
    except Exception:
        fallback = {"problem_statement": raw[:400], "research_question": "", "hypothesis": "", "objectives": [], "expected_contribution": ""}
        session["thesis"] = fallback
        return jsonify({"thesis": fallback})

@app.route("/api/thesis/save", methods=["POST"])
def save_thesis():
    data = request.get_json() or {}
    session["thesis"] = data.get("thesis", {})
    return jsonify({"success": True})

@app.route("/api/outline/generate", methods=["POST"])
def api_generate_outline():
    topic = session.get("selected_topic", "")
    ctx = session.get("project_context", {})
    thesis = session.get("thesis", {})
    if not topic:
        return jsonify({"error": "Complete previous structural steps first."}), 400

    thesis_str = thesis.get("problem_statement","") + " " + thesis.get("research_question","")

    prompt = f"""You are a paper structure expert. Create a detailed research paper outline.
Topic: {topic}
Thesis: {thesis_str}
Paper Length: {ctx.get('length','6-8 pages')}
Journal Format: {ctx.get('journal','IEEE')}

Return ONLY a JSON array of sections. Each section: title (string), subsections (list of strings).
Include: Abstract, Introduction, Literature Review, Methodology, Results, Discussion, Conclusion, References.
No markdown. No explanation."""

    raw = gemini_call(prompt)
    if is_error(raw):
        return jsonify({"error": raw}), 500
    try:
        clean = re.sub(r'```json|```', '', raw).strip()
        parsed = json.loads(clean)
        session["outline"] = parsed
        return jsonify({"outline": parsed})
    except Exception:
        fallback = [{"title": s.strip(), "subsections": []} for s in raw.split("\n") if s.strip()]
        session["outline"] = fallback
        return jsonify({"outline": fallback})

@app.route("/api/outline/save", methods=["POST"])
def save_outline():
    data = request.get_json() or {}
    session["outline"] = data.get("outline", [])
    return jsonify({"success": True})

@app.route("/api/sections/generate", methods=["POST"])
def api_generate_section():
    data = request.get_json() or {}
    sec = data.get("section", "")
    topic = session.get("selected_topic", "")
    ctx = session.get("project_context", {})
    thesis = session.get("thesis", {})
    refs = session.get("uploaded_references", [])
    exp = session.get("experiment_data", {})

    if not topic:
         return jsonify({"error": "Topic validation parameter missing."}), 400

    refs_text = " | ".join([f"{r.get('title','')} — {r.get('abstract','')[:300]}" for r in refs])
    thesis_str = thesis.get("problem_statement","") + " " + thesis.get("research_question","")

    if sec == "Introduction":
        prompt = f"""Write a professional Introduction section for a research paper.
Topic: {topic}
Thesis: {thesis_str}
Subject: {ctx.get('subject','')}
Academic Level: {ctx.get('level','')}

Write 3-4 paragraphs covering: background, motivation, research gap, paper organization.
Do NOT invent citations or statistics. Write [CITATION NEEDED] where a citation would normally appear.
Plain text only, no markdown headers."""
        content = gemini_call(prompt, 0.3)
    elif sec == "Literature Review":
        prompt = f"""Write a Literature Review section for a research paper.
Topic: {topic}
Available verified references: {refs_text[:2500] if refs_text else 'No references provided. Note this limitation.'}

Write 4-5 paragraphs analyzing existing work, identifying themes, and noting gaps.
ONLY cite from the provided references. Mark any additional citation as [UNVERIFIED - VERIFY BEFORE SUBMISSION].
Plain text only, no markdown."""
        content = gemini_call(prompt, 0.3)
    elif sec == "Methodology":
        exp_str = json.dumps(exp) if exp else "No experiment data provided."
        prompt = f"""Write a Methodology section for a research paper.
Topic: {topic}
Research Question: {thesis_str}
Experiment Details: {exp_str}

Write 3-4 paragraphs: research design, data collection, analysis approach, tools used.
Do NOT invent datasets, accuracy scores, or model names beyond what is provided.
If no experiment data, describe a general methodology framework.
Plain text only, no markdown headers."""
        content = gemini_call(prompt, 0.3)
    elif sec == "Discussion":
        results_note = session.get("paper_sections", {}).get("Results", "")
        prompt = f"""Write a Discussion section for a research paper.
Topic: {topic}
Thesis/Hypothesis: {thesis_str}
Results summary (student-provided): {results_note or 'Student has not entered results yet. Write a template with [RESULTS PLACEHOLDER] markers.'}

Write 3-4 paragraphs interpreting findings, comparing with literature, discussing implications, noting limitations.
Do NOT invent specific numbers or results. Use [RESULTS PLACEHOLDER] where actual data is needed.
Plain text only, no markdown headers."""
        content = gemini_call(prompt, 0.3)
    elif sec == "Conclusion":
        prompt = f"""Write a Conclusion section for a research paper.
Topic: {topic}
Thesis: {thesis_str}
Expected Contributions: {thesis.get('expected_contribution','')}

Write 2-3 paragraphs: summary of work, key findings, contributions, future work.
Do NOT invent results or statistics. Plain text only, no markdown headers."""
        content = gemini_call(prompt, 0.3)
    elif sec == "Abstract":
        other_secs = {k: v[:300] for k, v in session.get("paper_sections", {}).items()}
        prompt = f"Write a concise academic abstract (150-250 words) for a paper on: {topic}\nContext: {json.dumps(other_secs)}\nNo markdown."
        content = gemini_call(prompt, 0.3)
    elif sec == "Results":
        content = "[Enter your experimental results here. Use the Experiment Builder in Step 8 to generate a results summary.]"
    else:
        content = ""

    if not is_error(content):
        paper_sections = session.get("paper_sections", {})
        paper_sections[sec] = content
        session["paper_sections"] = paper_sections
        return jsonify({"success": True, "content": content})
    return jsonify({"error": content}), 500

@app.route("/api/sections/save", methods=["POST"])
def save_section_content():
    data = request.get_json() or {}
    sec = data.get("section", "")
    content = data.get("content", "")
    paper_sections = session.get("paper_sections", {})
    if sec:
        paper_sections[sec] = content
        session["paper_sections"] = paper_sections
        return jsonify({"success": True})
    return jsonify({"success": False}), 400

@app.route("/api/experiments/save", methods=["POST"])
def save_experiments():
    data = request.get_json() or {}
    session["experiment_data"] = data.get("experiment", {})
    return jsonify({"success": True})

@app.route("/api/experiments/summarize", methods=["POST"])
def api_summarize_experiments():
    exp = session.get("experiment_data", {})
    prompt = f"""Summarize these experiment results for a Results section of a research paper.
Experiment data: {json.dumps(exp)}

Write 2-3 paragraphs presenting the results clearly and objectively.
Only report the numbers provided. Do NOT invent additional results.
Use academic language. Plain text, no markdown."""
    
    summary = gemini_call(prompt, 0.2)
    if not is_error(summary):
        paper_sections = session.get("paper_sections", {})
        paper_sections["Results"] = summary
        session["paper_sections"] = paper_sections
        return jsonify({"success": True, "summary": summary})
    return jsonify({"error": summary}), 500

@app.route("/api/charts/generate", methods=["POST"])
def api_generate_chart():
    if "csv_file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file = request.files["csv_file"]
    chart_type = request.form.get("chart_type", "Bar")
    x_col = request.form.get("x_col", "")
    y_col = request.form.get("y_col", "")

    try:
        df = pd.read_csv(file)
        if x_col not in df.columns or y_col not in df.columns:
            return jsonify({"error": "Columns missing"}), 400
        
        plt.figure(figsize=(7, 4.5))
        if chart_type == "Bar":
            plt.bar(df[x_col].astype(str), df[y_col], color="#2563eb")
        elif chart_type == "Line":
            plt.plot(df[x_col], df[y_col], marker='o', color="#2563eb")
        elif chart_type == "Scatter":
            plt.scatter(df[x_col], df[y_col], color="#2563eb")
        elif chart_type == "Box":
            datasets = [group[y_col].values for name, group in df.groupby(x_col)]
            labels = [str(name) for name, group in df.groupby(x_col)]
            plt.boxplot(datasets, labels=labels)
        elif chart_type == "Pie":
            plt.pie(df[y_col], labels=df[x_col].astype(str), autopct='%1.1f%%')
            
        plt.title(f"{y_col} vs {x_col}", fontsize=12, pad=10)
        plt.xlabel(x_col)
        plt.ylabel(y_col)
        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, format="png", dpi=150)
        buf.seek(0)
        plt.close()

        return send_file(buf, mimetype="image/png", as_attachment=True, download_name="chart.png")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/citations/add_unverified", methods=["POST"])
def add_unverified_citation():
    data = request.get_json() or {}
    title = data.get("title", "")
    authors = data.get("authors", "")
    citations = session.get("citations", {"verified": [], "unverified": []})
    citations["unverified"].append({"title": title, "authors": authors})
    session["citations"] = citations
    return jsonify({"success": True})

@app.route("/api/revision/review", methods=["POST"])
def api_review_section():
    data = request.get_json() or {}
    sec_name = data.get("section", "")
    content = data.get("content", "")

    prompt = f"""Review this research paper section as an academic peer reviewer.
Section: {sec_name}
Content: {content[:2000]}

Return ONLY a JSON object with keys:
- overall_score (int 1-10)
- strengths (list of 2-3 strings)
- weaknesses (list of 2-3 strings)  
- suggestions (list of 3-4 actionable strings)
- citation_issues (list of strings, any unverified or missing citations)
No markdown. No explanation."""

    raw = gemini_call(prompt)
    if is_error(raw):
        return jsonify({"error": raw}), 500
    try:
        clean = re.sub(r'```json|```', '', raw).strip()
        parsed = json.loads(clean)
        review_results = session.get("review_results", {})
        review_results[f"review_{sec_name}"] = parsed
        session["review_results"] = review_results
        return jsonify({"review": parsed})
    except Exception:
        fallback = {"overall_score": 0, "strengths": [], "weaknesses": [raw[:200]], "suggestions": [], "citation_issues": []}
        return jsonify({"review": fallback})

@app.route("/api/revision/improve", methods=["POST"])
def api_improve_section():
    data = request.get_json() or {}
    sec_name = data.get("section", "")
    content = data.get("content", "")
    suggestions = data.get("suggestions", "")

    prompt = f"""Improve this research paper section based on the feedback. Preserve all factual claims and do not invent new information.
Section: {sec_name}
Feedback: {suggestions}
Original: {content[:2000]}

Return the improved section text only. No markdown headers. No explanation."""

    improved = gemini_call(prompt, 0.3)
    if not is_error(improved):
        review_results = session.get("review_results", {})
        review_results[f"improved_{sec_name}"] = improved
        session["review_results"] = review_results
        return jsonify({"improved": improved})
    return jsonify({"error": improved}), 500

@app.route("/api/formatting/guide", methods=["POST"])
def api_formatting_guide():
    data = request.get_json() or {}
    journal = data.get("journal", "IEEE")
    sec_name = data.get("section", "")
    content = data.get("content", "")

    prompt = f"""Provide formatting guidance for the {sec_name} section according to {journal} style.
Content preview: {content[:500]}

List 4-6 specific formatting requirements for {journal}: font, spacing, heading style, citation format, etc.
Be specific and actionable. Plain text, numbered list."""

    guidance = gemini_call(prompt, 0.2)
    return jsonify({"guidance": guidance})

@app.route("/api/export/docx", methods=["GET"])
def export_docx():
    if not DOCX_AVAILABLE:
        return "python-docx library not available", 500
    ctx = session.get("project_context", {})
    sections = session.get("paper_sections", {})
    
    doc = Document()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(ctx.get("title", "Research Paper"))
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph(ctx.get("student_names", ""), style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(ctx.get("college", ""), style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    for sec_name in ["Abstract", "Introduction", "Literature Review", "Methodology", "Results", "Discussion", "Conclusion"]:
        content = sections.get(sec_name, "")
        if content:
            doc.add_heading(sec_name, level=1)
            doc.add_paragraph(content)
            
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", as_attachment=True, download_name="research_paper.docx")

@app.route("/api/export/latex/generate", methods=["POST"])
def api_generate_latex():
    ctx = session.get("project_context", {})
    sections = session.get("paper_sections", {})
    refs = session.get("uploaded_references", [])
    thesis = session.get("thesis", {})

    title = ctx.get("title", "Research Paper")
    authors = ctx.get("student_names", "Author")
    institution = ctx.get("college", "University")

    main_tex = f"""\\documentclass[12pt,a4paper]{{article}}
\\usepackage[utf8]{{inputenc}}
\\usepackage{{amsmath,amssymb}}
\\usepackage{{graphicx}}
\\usepackage{{cite}}
\\usepackage{{hyperref}}
\\usepackage[margin=1in]{{geometry}}
\\usepackage{{times}}

\\title{{{title}}}
\\author{{{authors} \\\\ {institution}}}
\\date{{\\today}}

\\begin{{document}}
\\maketitle

\\begin{{abstract}}
{sections.get('Abstract', 'Abstract not yet generated.')}
\\end{{abstract}}

\\section{{Introduction}}
{sections.get('Introduction', 'Introduction not yet generated.')}

\\section{{Literature Review}}
{sections.get('Literature Review', 'Literature Review not yet generated.')}

\\section{{Methodology}}
{sections.get('Methodology', 'Methodology not yet generated.')}

\\section{{Results}}
{sections.get('Results', '[Results to be added after experiments are conducted.]')}

\\section{{Discussion}}
{sections.get('Discussion', 'Discussion not yet generated.')}

\\section{{Conclusion}}
{sections.get('Conclusion', 'Conclusion not yet generated.')}

\\bibliography{{references}}
\\bibliographystyle{{ieeetr}}
\\end{{document}}
"""

    bib = ""
    for i, ref in enumerate(refs, 1):
        key = f"ref{i}"
        bib += f"""@article{{{key},
  title = {{{ref.get('title','Unknown Title')}}},
  author = {{{ref.get('authors','Unknown Author')}}},
  note = {{Verified uploaded reference}},
  year = {{2024}}
}}\n\n"""

    meta = {
        "title": title,
        "authors": authors,
        "institution": institution,
        "subject": ctx.get("subject", ""),
        "citation_style": ctx.get("citation_style", "IEEE"),
        "thesis_statement": thesis.get("problem_statement", ""),
        "generated_by": "ResearchMentor Paper Assistant",
        "note": "Verify all citations before submission. Replace [CITATION NEEDED] and [UNVERIFIED] markers."
    }

    latex_pack = {"main.tex": main_tex, "references.bib": bib, "metadata.json": json.dumps(meta, indent=2)}
    session["generated_latex"] = latex_pack
    return jsonify({"success": True, "latex": latex_pack})

@app.route("/api/export/latex/download", methods=["GET"])
def download_latex_zip():
    latex = session.get("generated_latex", {})
    if not latex:
        return "No LaTeX generated yet", 400
        
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for fname, content in latex.items():
            zf.writestr(fname, content)
            
    zip_buf.seek(0)
    return send_file(zip_buf, mimetype="application/zip", as_attachment=True, download_name="research_paper_latex.zip")

if __name__ == "__main__":
    app.run(debug=True)
